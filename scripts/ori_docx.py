#!/usr/bin/env python3
"""Ori-native DOCX generator (see references/office.md).

Encodes Ori's token system for Word / Google Docs. Agents import OriDoc
and compose documents from the same components the HTML templates use.
Running this file directly writes ori-docx-demo.docx as a visual check.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# ── Ori tokens (mirror tokens/ori.css; change there first) ──────────
BLUE = RGBColor(0x2F, 0x80, 0xED)
BLUE_DEEP = RGBColor(0x11, 0x60, 0xC6)
NAVY = RGBColor(0x21, 0x3C, 0x60)
INK = RGBColor(0x00, 0x04, 0x16)
SLATE = RGBColor(0x4A, 0x4D, 0x5A)
MIST = RGBColor(0x8A, 0x94, 0xAB)
POS = RGBColor(0x0E, 0x8A, 0x5C)
NEG = RGBColor(0xD6, 0x45, 0x45)
LINE = "E3E8F4"
LINE_STRONG = "C9D2E8"
WASH = "EFF5FD"
WARN = "B7791F"

SANS = "Inter"
MONO = "IBM Plex Mono"

ASSETS = Path(__file__).resolve().parent.parent / "assets" / "logo" / "png"

MARGINS = {  # T, R, B, L in mm — mirrors design.md §4.2
    "one-pager": (14, 16, 14, 16),
    "long-doc": (18, 20, 20, 20),
    "report": (16, 18, 18, 18),
    "resume": (12, 14, 12, 14),
    "minutes": (16, 18, 18, 18),
}


# ── low-level oxml helpers ───────────────────────────────────────────
def _p_border(paragraph, edge="bottom", size=4, color=LINE, space=4):
    """Paragraph border. size is in 1/8 pt (4 = 0.5pt)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)


def _shade(cell, hex_fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_borders(cell, edges, size=4, color=LINE):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in edges:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _attach_tbl_borders(table, borders):
    """Insert w:tblBorders at its schema position: before tblLayout /
    tblCellMar / tblLook if present, else append. Word tolerates disorder;
    strict OOXML validators don't."""
    tblPr = table._tbl.tblPr
    for tag in ("w:tblBorders",):          # replace an existing block
        old = tblPr.find(qn(tag))
        if old is not None:
            tblPr.remove(old)
    for tag in ("w:tblLayout", "w:tblCellMar", "w:tblLook"):
        ref = tblPr.find(qn(tag))
        if ref is not None:
            ref.addprevious(borders)
            return
    tblPr.append(borders)


def _no_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    _attach_tbl_borders(table, borders)


def _outer_table_borders(table, size=4, color=LINE):
    """Hairline box around the table, no inner rules."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    for edge in ("insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        borders.append(el)
    _attach_tbl_borders(table, borders)


def _tracking(run, twentieths=12):
    """Letter-spacing in 1/20 pt (12 ≈ 0.6pt — mono label tracking)."""
    rPr = run._r.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(twentieths))
    rPr.append(sp)


def _field(paragraph, instr):
    """Insert a field code (PAGE / NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, end):
        run._r.append(el)
    return run


# ── the document builder ─────────────────────────────────────────────
class OriDoc:
    def __init__(self, doc_type="Document", doc_id="ST-DOC-000",
                 date="", classification="Internal", artifact="long-doc"):
        self.doc = Document()
        self.doc_id = doc_id
        sec = self.doc.sections[0]
        sec.page_width, sec.page_height = Mm(210), Mm(297)
        t, r, b, l = MARGINS.get(artifact, MARGINS["long-doc"])
        sec.top_margin, sec.right_margin = Mm(t), Mm(r)
        sec.bottom_margin, sec.left_margin = Mm(b), Mm(l)
        n = self.doc.styles["Normal"]
        n.font.name, n.font.size, n.font.color.rgb = SANS, Pt(10), INK
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing = 1.5
        self._usable = (sec.page_width - sec.left_margin
                        - sec.right_margin)
        self._meta_rail(doc_type, date, classification)
        self._footer()

    # ── chrome ──
    def _meta_rail(self, doc_type, date, classification):
        p = self.doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
        logo = ASSETS / "mark.png"
        if logo.exists():
            p.add_run().add_picture(str(logo), height=Pt(12))
            self._mono(p, "  ")
        self._mono(p, "SOCIALTRAIT", color=INK, bold=True)
        parts = [x for x in (doc_type, date) if x]
        self._mono(p, "  ·  " + "  ·  ".join(x.upper() for x in parts))
        self._mono(p, "\t" + classification.upper())
        _p_border(p, "bottom", size=4, color=LINE, space=6)
        p.paragraph_format.space_after = Pt(14)

    def _footer(self):
        f = self.doc.sections[0].footer
        p = f.paragraphs[0]
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
        self._mono(p, "▪ ", color=BLUE, size=7.5)
        self._mono(p, f"SOCIALTRAIT / ORI · {self.doc_id}\t", size=7.5)
        self._mono(p, "PAGE ", size=7.5)
        r = _field(p, "PAGE")
        self._style_run(r, MONO, 7.5, MIST)
        self._mono(p, " / ", size=7.5)
        r = _field(p, "NUMPAGES")
        self._style_run(r, MONO, 7.5, MIST)
        _p_border(p, "top", size=4, color=LINE, space=6)

    # ── run helpers ──
    @staticmethod
    def _style_run(run, font, size, color, bold=False, caps=False):
        run.font.name, run.font.size = font, Pt(size)
        run.font.color.rgb, run.font.bold = color, bold
        if caps:
            run.font.all_caps = True
        return run

    def _mono(self, p, text, color=MIST, size=8, bold=False):
        r = p.add_run(text)
        self._style_run(r, MONO, size, color, bold=bold, caps=True)
        _tracking(r, 12)
        return r

    def _rich(self, p, text, size, color, bold_color=INK):
        """Minimal markup: <b>bold</b> renders weight-600 ink."""
        for i, part in enumerate(text.split("<b>")):
            if i == 0:
                if part:
                    self._style_run(p.add_run(part), SANS, size, color)
                continue
            bold_txt, _, rest = part.partition("</b>")
            self._style_run(p.add_run(bold_txt), SANS, size, bold_color,
                            bold=True)
            if rest:
                self._style_run(p.add_run(rest), SANS, size, color)

    # ── components ──
    def title(self, text):
        p = self.doc.add_paragraph()
        r = self._style_run(p.add_run(text), SANS, 22, INK, bold=True)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        return p

    def lede(self, text):
        p = self.doc.add_paragraph()
        self._style_run(p.add_run(text), SANS, 11, SLATE)
        p.paragraph_format.space_after = Pt(14)

    def section(self, idx, eyebrow, heading):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(2)
        self._mono(p, f"{idx:02d}  ", color=MIST)
        self._mono(p, eyebrow, color=BLUE)
        h = self.doc.add_paragraph()
        self._style_run(h.add_run(heading), SANS, 15, INK, bold=True)
        h.paragraph_format.line_spacing = 1.25
        h.paragraph_format.space_after = Pt(4)
        _p_border(h, "bottom", size=8, color=LINE, space=4)  # thread (gray; blue lives in eyebrow)
        h.paragraph_format.space_after = Pt(10)

    def h3(self, text):
        p = self.doc.add_paragraph()
        self._style_run(p.add_run(text), SANS, 11.5, INK, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)

    def body(self, text):
        p = self.doc.add_paragraph()
        self._rich(p, text, 10, INK)
        return p

    def bullets(self, items):
        for it in items:
            p = self.doc.add_paragraph(style="List Bullet")
            self._rich(p, it, 10, INK)
            p.paragraph_format.space_after = Pt(3)

    @staticmethod
    def _fixed_layout(table):
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        table._tbl.tblPr.append(layout)

    def metric_row(self, metrics):
        """metrics: [{label, value, note?, hero?}] — 3 or 4 of them."""
        t = self.doc.add_table(rows=3, cols=len(metrics))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False
        self._fixed_layout(t)
        _no_table_borders(t)
        col_w = int(self._usable / len(metrics))
        for row in t.rows:
            for cell in row.cells:
                cell.width = col_w
        for i, m in enumerate(metrics):
            for row in (0, 1, 2):
                cell = t.cell(row, i)
                edges = ["left", "right"] if True else []
                if row == 0:
                    edges.append("top")
                if row == 2:
                    edges.append("bottom")
                _cell_borders(cell, edges, size=4, color=LINE)
            lp = t.cell(0, i).paragraphs[0]
            self._mono(lp, m["label"], color=MIST)
            lp.paragraph_format.space_before = Pt(6)
            vp = t.cell(1, i).paragraphs[0]
            r = vp.add_run(m["value"])
            self._style_run(r, MONO, 24, BLUE if m.get("hero") else INK,
                            bold=True)
            np_ = t.cell(2, i).paragraphs[0]
            self._style_run(np_.add_run(m.get("note", "")), SANS, 8.5, SLATE)
            np_.paragraph_format.space_after = Pt(6)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def callout(self, label, text, kind="insight"):
        color = {"insight": BLUE, "warn": RGBColor(0xB7, 0x79, 0x1F),
                 "risk": NEG}[kind]
        border = {"insight": "2F80ED", "warn": WARN, "risk": "D64545"}[kind]
        t = self.doc.add_table(rows=1, cols=1)
        _no_table_borders(t)
        cell = t.cell(0, 0)
        _shade(cell, WASH)
        _cell_borders(cell, ["left"], size=18, color=border)
        lp = cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(5)
        self._mono(lp, label, color=color, size=7.5)
        bp = cell.add_paragraph()
        self._rich(bp, text, 10, INK)
        bp.paragraph_format.space_after = Pt(5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def table(self, headers, rows, data_cols=()):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        _no_table_borders(t)
        for j, htext in enumerate(headers):
            cell = t.cell(0, j)
            _cell_borders(cell, ["bottom"], size=6, color=LINE_STRONG)
            p = cell.paragraphs[0]
            self._mono(p, htext, color=MIST)
            if j in data_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.cell(1 + i, j)
                _cell_borders(cell, ["bottom"], size=4, color=LINE)
                p = cell.paragraphs[0]
                if j in data_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    self._style_run(p.add_run(str(val)), MONO, 9, INK)
                else:
                    self._style_run(p.add_run(str(val)), SANS, 9.5, INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def quote(self, text, cite):
        t = self.doc.add_table(rows=1, cols=1)
        _no_table_borders(t)
        cell = t.cell(0, 0)
        _cell_borders(cell, ["left"], size=12, color=LINE_STRONG)
        p = cell.paragraphs[0]
        self._style_run(p.add_run(f"“{text}”"), SANS, 11, INK)
        cp = cell.add_paragraph()
        self._mono(cp, "— ", color=BLUE)
        self._mono(cp, cite, color=MIST)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def kv(self, pairs):
        t = self.doc.add_table(rows=len(pairs), cols=2)
        _no_table_borders(t)
        t.columns[0].width = Cm(3.6)
        for i, (k, v) in enumerate(pairs):
            self._mono(t.cell(i, 0).paragraphs[0], k, color=MIST)
            self._style_run(t.cell(i, 1).paragraphs[0].add_run(v),
                            SANS, 9.5, INK)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def logistics(self, left, right):
        """Meeting logistics: two kv columns in a hairline box."""
        rows = max(len(left), len(right))
        t = self.doc.add_table(rows=rows, cols=4)
        t.autofit = False
        self._fixed_layout(t)
        _outer_table_borders(t, size=4, color=LINE)
        widths = [0.13, 0.37, 0.13, 0.37]
        for r in t.rows:
            for j, cell in enumerate(r.cells):
                cell.width = int(self._usable * widths[j])
        for col, pairs in ((0, left), (2, right)):
            for i, (k, v) in enumerate(pairs):
                kp = t.cell(i, col).paragraphs[0]
                self._mono(kp, k, color=MIST)
                vp = t.cell(i, col + 1).paragraphs[0]
                self._style_run(vp.add_run(v), SANS, 9.5, INK)
                if i == 0:
                    kp.paragraph_format.space_before = Pt(5)
                    vp.paragraph_format.space_before = Pt(5)
                if i == rows - 1:
                    kp.paragraph_format.space_after = Pt(5)
                    vp.paragraph_format.space_after = Pt(5)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def agenda(self, items):
        """Pre-meeting agenda: eyebrow + mono-indexed compact list."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        self._mono(p, "Agenda", color=BLUE)
        _p_border(p, "bottom", size=8, color=LINE, space=4)
        for i, item in enumerate(items, 1):
            ip = self.doc.add_paragraph()
            self._mono(ip, f"{i:02d}  ", color=BLUE, size=8.5)
            self._rich(ip, item, 9.5, INK)
            if i < len(items):    # template suppresses the last hairline
                _p_border(ip, "bottom", size=4, color=LINE, space=3)
            ip.paragraph_format.space_after = Pt(4)

    def decisions(self, items):
        """Meeting decisions: node glyph + referenceable D-ids."""
        for i, item in enumerate(items, 1):
            p = self.doc.add_paragraph()
            self._mono(p, "\u25aa ", color=BLUE, size=8.5)
            self._mono(p, f"D{i}  ", color=BLUE, size=8.5)
            self._rich(p, item, 10, INK)
            _p_border(p, "bottom", size=4, color=LINE, space=3)
            p.paragraph_format.space_after = Pt(4)

    def actions(self, rows):
        """Action items: [{action, owner, due, status}] -> A-id table.
        status in {open, done, blocked} -> blue / green / red mono."""
        status_color = {"open": BLUE_DEEP, "done": POS, "blocked": NEG}
        t = self.doc.add_table(rows=1 + len(rows), cols=5)
        t.autofit = False
        self._fixed_layout(t)
        _no_table_borders(t)
        widths = [0.07, 0.47, 0.17, 0.12, 0.17]
        for r in t.rows:
            for j, cell in enumerate(r.cells):
                cell.width = int(self._usable * widths[j])
        for j, htext in enumerate(["#", "Action", "Owner", "Due", "Status"]):
            cell = t.cell(0, j)
            _cell_borders(cell, ["bottom"], size=6, color=LINE_STRONG)
            self._mono(cell.paragraphs[0], htext, color=MIST)
        for i, row in enumerate(rows, 1):
            cells = t.rows[i].cells
            for cell in cells:
                _cell_borders(cell, ["bottom"], size=4, color=LINE)
            self._mono(cells[0].paragraphs[0], f"A{i}", color=BLUE, size=8.5)
            self._rich(cells[1].paragraphs[0], row["action"], 9.5, INK)
            self._mono(cells[2].paragraphs[0], row["owner"], color=SLATE,
                       size=8.5)
            self._mono(cells[3].paragraphs[0], row.get("due", "TBD"),
                       color=SLATE, size=8.5)
            st = row.get("status", "open").lower()
            self._mono(cells[4].paragraphs[0], st,
                       color=status_color.get(st, BLUE_DEEP), size=8)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def smallprint(self, text):
        """Next meeting / status / distribution / confidentiality."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        self._rich(p, text, 8, MIST, bold_color=SLATE)
        return p

    def page_break(self):
        self.doc.add_page_break()

    def save(self, path):
        self.doc.save(path)
        return path


# ── demo / visual regression check ──────────────────────────────────
if __name__ == "__main__":
    d = OriDoc(doc_type="Statement of work", doc_id="ST-SOW-007",
               date="2026-07-08", classification="Client confidential")
    d.title("Audience simulation pilot — statement of work")
    d.lede("Scope, timeline, and commercial terms for the Q3 simulated-"
           "audience pilot between Socialtrait and Arlo Foods.")
    d.metric_row([
        {"label": "Studies", "value": "8", "note": "two campaigns"},
        {"label": "Duration", "value": "10", "note": "weeks, from kickoff",
         "hero": True},
        {"label": "Fee", "value": "$44k", "note": "fixed, net-30"},
        {"label": "Panel size", "value": "2,400", "note": "per study"},
    ])
    d.section(1, "Scope", "Six simulated studies across two campaigns")
    d.body("Socialtrait will run <b>eight simulated-audience studies</b> "
           "across the client's two Q3 campaigns, covering four audience "
           "cells per study with calibrated persona models.")
    d.bullets(["Creative pre-tests for 6 variants per campaign",
               "Segment-level preference ranking with confidence intervals",
               "Verbatim-style persona reactions for creative iteration"])
    d.section(2, "Timeline", "Kickoff to final readout in ten weeks")
    d.table(["Phase", "Deliverable", "Week"],
            [["Onboard", "Data intake, persona calibration", "1–2"],
             ["Studies", "8 studies, rolling readouts", "3–8"],
             ["Validation", "Holdout comparison, final report", "9–10"]],
            data_cols=(2,))
    d.quote("The simulation caught in four hours what our panel would have "
            "told us three weeks after launch.", "VP Growth · CPG pilot")
    d.callout("The ask", "Countersign by <b>July 18</b> so persona "
              "calibration completes before the August media flight.")
    print("wrote", d.save(str(Path.cwd() / "ori-docx-demo.docx")))
